"""MinerU 智能文档解析模块

基于 MinerU 引擎实现复杂文档的智能解析：
1. PDF 文档：版面分析、表格识别、公式识别
2. Word 文档：段落提取、表格转换
3. 输出：Markdown 结构化文本

MinerU 特性：
- 版面恢复：保留原始文档的层级结构
- 表格识别：自动识别并转换为 Markdown 表格
- 公式处理：LaTeX 公式保留
- OCR 识别：支持扫描件 PDF

硬件要求：
- 推荐 GPU: NVIDIA GPU with 4GB+ VRAM
- 本项目使用 RTX 4050 (8GB) 实测可用
- CPU 模式可用，但速度较慢

依赖安装：
    pip install magic-pdf[full] -i https://pypi.tuna.tsinghua.edu.cn/simple

使用示例：
    from app.services.document_parser import MinerUParser
    
    parser = MinerUParser()
    result = parser.parse("document.pdf")
    print(result.content)  # Markdown 格式内容
"""

import os
import sys
from pathlib import Path
from typing import Dict, List, Optional, Any, Union
from dataclasses import dataclass, field

from loguru import logger

# 延迟导入 MinerU 相关模块（magic-pdf v1.3.12 API）
_magic_pdf_available = False
_magic_pdf_import_error = None

try:
    from magic_pdf.tools.common import do_parse
    from magic_pdf.data.read_api import PymuDocDataset
    from magic_pdf.libs.config_reader import get_device
    import magic_pdf.model as model_config
    _magic_pdf_available = True
except ImportError as e:
    _magic_pdf_import_error = str(e)
    logger.warning(f"MinerU (magic-pdf) 未安装: {e}")
    logger.warning("将使用降级解析方案 (PyMuPDF + python-docx)")


@dataclass
class MinerUConfig:
    """MinerU 配置类
    
    用于配置 MinerU 解析引擎的行为参数。
    """
    # 解析模式
    parse_mode: str = "auto"  # auto: 自动选择, ocr: 强制 OCR, txt: 仅文本
    
    # 表格识别
    table_mode: str = "smart"  # smart: 智能识别, html: 输出 HTML, markdown: 输出 Markdown
    
    # 公式处理
    formula_mode: str = "latex"  # latex: LaTeX 格式, unicode: Unicode 格式
    
    # OCR 配置
    ocr_enabled: bool = True  # 是否启用 OCR
    ocr_lang: str = "ch"  # 中文: ch, 英文: en, 中英混合: auto
    
    # 输出配置
    output_formats: List[str] = field(default_factory=lambda: ["text", "table", "formula"])
    
    # GPU 配置
    gpu_device: int = 0  # GPU 设备号，-1 表示使用 CPU
    
    # 推理批处理大小
    batch_size: int = 1
    
    def to_dict(self) -> Dict[str, Any]:
        return {
            "parse_mode": self.parse_mode,
            "table_mode": self.table_mode,
            "formula_mode": self.formula_mode,
            "ocr_enabled": self.ocr_enabled,
            "ocr_lang": self.ocr_lang,
            "output_formats": self.output_formats,
            "gpu_device": self.gpu_device,
            "batch_size": self.batch_size,
        }


class MinerUParser:
    """MinerU 智能文档解析器
    
    支持 PDF、Word 等复杂文档的智能解析，
    自动进行版面分析、表格识别和结构化转换。
    
    特性：
    1. PDF 解析
       - 文本提取：精准提取 PDF 中的文字内容
       - 表格识别：自动识别并转换为 Markdown 表格
       - 公式识别：支持数学公式的 LaTeX 转换
       - 版面分析：保留文档的层级结构
       - OCR 识别：支持扫描件和图片类 PDF
    
    2. Word 解析
       - 段落提取：保留段落样式和层级
       - 表格转换：自动转换为 Markdown 表格
       - 标题识别：识别并标记文档结构
    
    3. 输出格式
       - Markdown：结构化输出，便于 RAG 知识库构建
       - 元数据：包含页码、位置等详细信息
    
    使用示例：
        parser = MinerUParser()
        result = parser.parse("document.pdf")
        print(result.content)
        print(result.metadata)
    """
    
    # 支持的文件类型
    SUPPORTED_EXTENSIONS = ['.pdf', '.docx', '.doc']
    
    def __init__(self, config: Optional[MinerUConfig] = None):
        """初始化 MinerU 解析器
        
        Args:
            config: MinerU 配置对象
        """
        self.config = config or MinerUConfig()
        self._initialized = False
        self._parser_instance = None
        self._ocr_processor = None
        
        # 检测 GPU
        self._check_gpu()
        
        # 初始化解析器
        self._initialize_parser()
    
    def _check_gpu(self) -> None:
        """检查 GPU 可用性"""
        try:
            import torch
            if torch.cuda.is_available():
                gpu_name = torch.cuda.get_device_name(0)
                gpu_memory = torch.cuda.get_device_properties(0).total_memory / 1024**3
                logger.info(f"检测到 GPU: {gpu_name}, 显存: {gpu_memory:.1f} GB")
                self._has_gpu = True
            else:
                logger.info("未检测到 GPU，将使用 CPU 模式")
                self._has_gpu = False
        except ImportError:
            logger.info("PyTorch 未安装，将使用 CPU 模式")
            self._has_gpu = False
    
    def _initialize_parser(self) -> None:
        """初始化 MinerU 解析器"""
        if not _magic_pdf_available:
            logger.warning(
                "MinerU (magic-pdf) 未安装，"
                "将使用降级方案 (PyMuPDF + python-docx)。"
                f"安装命令: pip install magic-pdf[full] -i https://pypi.tuna.tsinghua.edu.cn/simple"
            )
            self._initialized = False
            return

        try:
            # 设置设备
            if self._has_gpu:
                os.environ['CUDA_VISIBLE_DEVICES'] = str(self.config.gpu_device)

            # 检查 MinerU 配置
            device = get_device()
            logger.info(f"MinerU 设备模式: {device}")

            self._initialized = True
            logger.info("MinerU 解析器初始化完成")

        except Exception as e:
            logger.error(f"MinerU 初始化失败: {e}")
            self._initialized = False
    
    def supports(self, file_path: str) -> bool:
        """判断是否支持该文件类型
        
        Args:
            file_path: 文件路径
            
        Returns:
            bool: 是否支持
        """
        ext = Path(file_path).suffix.lower()
        return ext in self.SUPPORTED_EXTENSIONS
    
    def parse(self, file_path: str) -> 'ParseResult':
        """解析文档
        
        Args:
            file_path: 文档路径
            
        Returns:
            ParseResult: 解析结果，包含 Markdown 内容和元数据
            
        Raises:
            FileNotFoundError: 文件不存在
            ValueError: 文件格式不支持
            RuntimeError: 解析过程出错
        """
        from app.services.document_parser.base_parser import ParseResult, BaseParser
        
        path = Path(file_path)
        
        # 验证文件
        if not path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")
        
        if not self.supports(file_path):
            raise ValueError(f"MinerUParser 不支持文件类型: {path.suffix}")
        
        warnings = []
        ext = path.suffix.lower()
        
        try:
            if ext == '.pdf':
                return self._parse_pdf(file_path, warnings)
            elif ext in ['.docx', '.doc']:
                return self._parse_docx(file_path, warnings)
            else:
                raise ValueError(f"不支持的文件类型: {ext}")
                
        except FileNotFoundError:
            raise
        except ValueError:
            raise
        except Exception as e:
            logger.error(f"文档解析失败: {e}")
            raise RuntimeError(f"文档解析失败: {e}") from e
    
    def _parse_pdf(self, file_path: str, warnings: List[str]) -> 'ParseResult':
        """解析 PDF 文档
        
        Args:
            file_path: PDF 文件路径
            warnings: 警告信息列表
            
        Returns:
            ParseResult: 解析结果
        """
        from app.services.document_parser.base_parser import ParseResult
        
        metadata = {
            "file_name": Path(file_path).name,
            "file_size": os.path.getsize(file_path),
            "file_type": "pdf",
            "parser": "MinerUParser",
            "gpu_used": self._has_gpu,
            "parse_mode": self.config.parse_mode,
        }
        
        # 根据 MinerU 是否可用选择解析方法
        if _magic_pdf_available and self._initialized:
            return self._parse_pdf_with_mineru(file_path, warnings, metadata)
        else:
            return self._parse_pdf_fallback(file_path, warnings, metadata)
    
    def _parse_pdf_with_mineru(
        self,
        file_path: str,
        warnings: List[str],
        metadata: Dict[str, Any]
    ) -> 'ParseResult':
        """使用 MinerU 解析 PDF

        Args:
            file_path: PDF 文件路径
            warnings: 警告信息列表
            metadata: 元数据字典

        Returns:
            ParseResult: 解析结果
        """
        from app.services.document_parser.base_parser import ParseResult

        try:
            logger.info(f"使用 MinerU 解析 PDF: {file_path}")

            pdf_path = Path(file_path)
            pdf_bytes = pdf_path.read_bytes()
            pdf_file_name = pdf_path.stem

            # 创建临时输出目录
            import tempfile
            output_dir = tempfile.mkdtemp(prefix="mineru_")

            # 映射解析模式
            parse_method_map = {
                "auto": "auto",
                "ocr": "ocr",
                "txt": "txt",
            }
            parse_method = parse_method_map.get(
                self.config.parse_mode, "auto"
            )

            # 调用 MinerU do_parse（返回 Markdown 和 JSON 文件路径）
            do_parse(
                output_dir=output_dir,
                pdf_file_name=pdf_file_name,
                pdf_bytes_or_dataset=pdf_bytes,
                model_list=[],  # 使用内置模型
                parse_method=parse_method,
                f_draw_span_bbox=False,
                f_draw_layout_bbox=False,
                f_dump_md=True,
                f_dump_middle_json=True,
                f_dump_model_json=False,
                f_dump_orig_pdf=False,
                f_dump_content_list=False,
                lang=self.config.ocr_lang if self.config.ocr_enabled else None,
            )

            # 读取生成的 Markdown 文件
            md_path = Path(output_dir) / f"{pdf_file_name}.md"
            if md_path.exists():
                content = md_path.read_text(encoding="utf-8")
            else:
                # 尝试从 content_list 重建
                content = self._rebuild_from_content_list(output_dir, pdf_file_name)

            # 读取 middle_json 获取页数等信息
            import json
            middle_json_path = Path(output_dir) / f"{pdf_file_name}_middle.json"
            table_count = 0
            formula_count = 0
            page_count = 0
            if middle_json_path.exists():
                try:
                    with open(middle_json_path, "r", encoding="utf-8") as f:
                        middle_data = json.load(f)
                    if isinstance(middle_data, list):
                        page_count = len(middle_data)
                    elif isinstance(middle_data, dict):
                        page_count = middle_data.get("page_count", len(middle_data.get("pdf_info", [])))
                        pdf_info = middle_data.get("pdf_info", [])
                        for page in pdf_info:
                            for block in page.get("preproc_blocks", []):
                                if block.get("type") == "table":
                                    table_count += 1
                except Exception:
                    pass

            metadata.update({
                "page_count": page_count,
                "table_count": table_count,
                "formula_count": formula_count,
                "content_length": len(content),
                "output_dir": output_dir,
            })

            content = self._sanitize_content(content)

            logger.info(
                f"MinerU PDF 解析完成: {page_count} 页, "
                f"{table_count} 个表格, "
                f"{len(content)} 字符"
            )

            return ParseResult(
                content=content,
                metadata=metadata,
                warnings=warnings
            )

        except Exception as e:
            logger.error(f"MinerU PDF 解析失败: {e}，尝试降级解析")
            warnings.append(f"MinerU 解析失败，使用降级方案: {e}")
            return self._parse_pdf_fallback(file_path, warnings, metadata)

    def _rebuild_from_content_list(
        self, output_dir: str, pdf_file_name: str
    ) -> str:
        """从 content_list JSON 重建 Markdown 内容"""
        content_list_path = Path(output_dir) / f"{pdf_file_name}_content_list.json"
        if not content_list_path.exists():
            return ""

        import json
        try:
            with open(content_list_path, "r", encoding="utf-8") as f:
                content_list = json.load(f)

            parts = []
            for item in content_list:
                block_type = item.get("type", "text")
                text = item.get("text", "")

                if block_type == "text":
                    parts.append(text)
                elif block_type == "title":
                    parts.append(f"## {text}")
                elif block_type == "table":
                    table_md = self._convert_table_to_markdown(item.get("table_data", []))
                    if table_md:
                        parts.append(table_md)
                elif block_type == "formula":
                    parts.append(f"${item.get('content', '')}$")

            return "\n\n".join(parts)
        except Exception:
            return ""
    
    def get_status(self) -> Dict[str, Any]:
        """获取解析器状态

        Returns:
            Dict: 状态信息
        """
        return {
            "available": self.is_available(),
            "gpu_available": self._has_gpu,
            "magic_pdf_available": _magic_pdf_available,
            "magic_pdf_error": _magic_pdf_import_error,
            "config": self.config.to_dict(),
        }

    def _extract_page_content(self, page_info: Dict) -> str:
        """从页面信息中提取 Markdown 内容
        
        Args:
            page_info: 页面解析结果
            
        Returns:
            str: Markdown 格式的页面内容
        """
        parts = []
        
        for block in page_info.get('content_blocks', []):
            block_type = block.get('type', 'text')
            
            if block_type == 'text':
                # 文本块
                text = block.get('content', '').strip()
                if text:
                    # 检查是否是标题
                    if block.get('is_title'):
                        level = block.get('title_level', 1)
                        parts.append(f"{'#' * level} {text}")
                    else:
                        parts.append(text)
                        
            elif block_type == 'table':
                # 表格块
                table_md = self._convert_table_to_markdown(block.get('table_data', []))
                if table_md:
                    parts.append(table_md)
                    
            elif block_type == 'formula':
                # 公式块
                formula = block.get('content', '')
                if formula:
                    if self.config.formula_mode == 'latex':
                        parts.append(f"${formula}$")
                    else:
                        parts.append(formula)
            
            elif block_type == 'image':
                # 图片块（记录图片位置）
                image_desc = block.get('description', '[图片]')
                parts.append(f"![{image_desc}](image_at_page_{block.get('page_num', 0)})")
        
        return '\n\n'.join(parts)
    
    def _convert_table_to_markdown(self, table_data: List[List[str]]) -> str:
        """将表格数据转换为 Markdown 格式
        
        Args:
            table_data: 表格数据，二维数组
            
        Returns:
            str: Markdown 表格格式
        """
        if not table_data or not table_data[0]:
            return ""
        
        lines = []
        for i, row in enumerate(table_data):
            # 转义 Markdown 特殊字符
            escaped_row = [
                cell.replace('|', '\\|').replace('\n', ' ')
                for cell in row
            ]
            lines.append('| ' + ' | '.join(escaped_row) + ' |')
            
            # 表头后添加分隔行
            if i == 0:
                separator = '| ' + ' | '.join(['---'] * len(row)) + ' |'
                lines.append(separator)
        
        return '\n'.join(lines)
    
    def _parse_pdf_fallback(
        self, 
        file_path: str, 
        warnings: List[str],
        metadata: Dict[str, Any]
    ) -> 'ParseResult':
        """降级解析 PDF（使用 PyMuPDF）
        
        当 MinerU 不可用时使用 PyMuPDF 进行基础解析。
        
        Args:
            file_path: PDF 文件路径
            warnings: 警告信息列表
            metadata: 元数据字典
            
        Returns:
            ParseResult: 解析结果
        """
        from app.services.document_parser.base_parser import ParseResult
        
        warnings.append("使用 PyMuPDF 降级解析，部分功能可能不可用")
        
        try:
            import fitz  # PyMuPDF
            
            doc = fitz.open(file_path)
            page_count = len(doc)
            
            content_parts = []
            for page_num in range(page_count):
                page = doc[page_num]
                
                # 提取文本
                text = page.get_text("text")
                
                # 如果文本为空且启用了 OCR，尝试 OCR
                if not text.strip() and self.config.ocr_enabled:
                    # 尝试从图片提取文本
                    pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
                    img_bytes = pix.tobytes("png")
                    ocr_text = self._ocr_image(img_bytes)
                    if ocr_text:
                        text = ocr_text
                        warnings.append(f"第 {page_num + 1} 页使用了 OCR")
                
                if text.strip():
                    content_parts.append(f"## 第 {page_num + 1} 页\n\n{text.strip()}")
            
            content = '\n\n'.join(content_parts)
            
            # 尝试识别表格（简单方法：查找由分隔符组成的行）
            content = self._extract_tables_from_text(content)
            
            metadata.update({
                "page_count": page_count,
                "table_count": content.count('|'),
                "formula_count": 0,
                "content_length": len(content),
                "parse_method": "PyMuPDF",
            })
            
            content = self._sanitize_content(content)
            
            logger.info(
                f"PDF 降级解析完成: {page_count} 页, "
                f"{len(content)} 字符"
            )
            
            return ParseResult(
                content=content,
                metadata=metadata,
                warnings=warnings
            )
            
        except ImportError:
            warnings.append("PyMuPDF 也未安装，尝试基础方法")
            return self._parse_pdf_basic(file_path, warnings, metadata)
        except Exception as e:
            warnings.append(f"PDF 解析出错: {e}")
            return self._parse_pdf_basic(file_path, warnings, metadata)
    
    def _parse_pdf_basic(
        self, 
        file_path: str, 
        warnings: List[str],
        metadata: Dict[str, Any]
    ) -> 'ParseResult':
        """最基础的 PDF 解析方法（使用标准库）
        
        Args:
            file_path: PDF 文件路径
            warnings: 警告信息列表
            metadata: 元数据字典
            
        Returns:
            ParseResult: 解析结果
        """
        from app.services.document_parser.base_parser import ParseResult
        
        warnings.append("使用基础文本提取，可能丢失格式信息")
        
        try:
            # 尝试使用 pdfplumber
            import pdfplumber
            
            with pdfplumber.open(file_path) as pdf:
                page_count = len(pdf.pages)
                
                content_parts = []
                for i, page in enumerate(pdf.pages):
                    text = page.extract_text() or ""
                    if text.strip():
                        content_parts.append(f"## 第 {i + 1} 页\n\n{text.strip()}")
                    
                    # 尝试提取表格
                    tables = page.extract_tables()
                    for table in tables:
                        table_md = self._convert_table_to_markdown(table)
                        content_parts.append(table_md)
                
                content = '\n\n'.join(content_parts)
                
                metadata.update({
                    "page_count": page_count,
                    "content_length": len(content),
                    "parse_method": "pdfplumber",
                })
                
        except ImportError:
            warnings.append("所有 PDF 解析库都不可用")
            content = ""
            metadata.update({
                "page_count": 0,
                "content_length": 0,
                "parse_method": "none",
            })
        
        return ParseResult(
            content=self._sanitize_content(content),
            metadata=metadata,
            warnings=warnings
        )
    
    def _ocr_image(self, image_bytes: bytes) -> Optional[str]:
        """对图片进行 OCR 识别
        
        Args:
            image_bytes: 图片字节数据
            
        Returns:
            Optional[str]: 识别的文本，如果失败返回 None
        """
        try:
            # 尝试使用 pytesseract
            import pytesseract
            from PIL import Image
            import io
            
            image = Image.open(io.BytesIO(image_bytes))
            text = pytesseract.image_to_string(image, lang=self.config.ocr_lang)
            return text.strip() if text else None
            
        except ImportError:
            logger.debug("pytesseract 未安装，无法进行 OCR")
            return None
        except Exception as e:
            logger.debug(f"OCR 识别失败: {e}")
            return None
    
    def _extract_tables_from_text(self, text: str) -> str:
        """从文本中提取并转换表格
        
        简单实现：查找由分隔符组成的行并转换为 Markdown 表格。
        
        Args:
            text: 原始文本
            
        Returns:
            str: 处理后的文本
        """
        lines = text.split('\n')
        result_lines = []
        in_table = False
        table_rows = []
        
        for line in lines:
            # 检测是否是表格分隔符行（如 |---|---|）
            stripped = line.strip()
            if stripped.startswith('|') and all(
                part.strip() in ('---', ':--', '--:', ':-:') 
                for part in stripped.split('|')[1:-1] 
                if part.strip()
            ):
                # 是分隔符行
                if table_rows:
                    table_rows.append(line)
                in_table = True
            elif in_table:
                # 在表格中
                if stripped.startswith('|'):
                    table_rows.append(line)
                else:
                    # 表格结束，输出 Markdown 表格
                    if table_rows:
                        result_lines.extend(table_rows)
                        result_lines.append('')  # 空行分隔
                    in_table = False
                    table_rows = []
                    result_lines.append(line)
            else:
                # 不在表格中
                if table_rows:
                    result_lines.extend(table_rows)
                    result_lines.append('')
                    table_rows = []
                result_lines.append(line)
        
        # 处理末尾可能的表格
        if table_rows:
            result_lines.extend(table_rows)
        
        return '\n'.join(result_lines)
    
    def _parse_docx(self, file_path: str, warnings: List[str]) -> 'ParseResult':
        """解析 Word 文档
        
        Args:
            file_path: Word 文件路径
            warnings: 警告信息列表
            
        Returns:
            ParseResult: 解析结果
        """
        from app.services.document_parser.base_parser import ParseResult
        
        metadata = {
            "file_name": Path(file_path).name,
            "file_size": os.path.getsize(file_path),
            "file_type": "docx",
            "parser": "MinerUParser",
        }
        
        try:
            from docx import Document
            
            doc = Document(file_path)
            
            content_parts = []
            current_heading = ""
            table_count = 0
            
            for element in doc.element.body:
                # 处理段落
                if element.tag.endswith('p'):
                    para = None
                    for p in doc.paragraphs:
                        if p._element == element:
                            para = p
                            break
                    
                    if para:
                        text = para.text.strip()
                        if text:
                            # 判断是否是标题
                            if para.style.name.startswith('Heading'):
                                level = int(para.style.name.replace('Heading ', ''))
                                content_parts.append(f"{'#' * level} {text}")
                                current_heading = text
                            else:
                                content_parts.append(text)
                
                # 处理表格
                elif element.tag.endswith('tbl'):
                    # 找到对应的表格
                    for table in doc.tables:
                        if table._element == element:
                            table_count += 1
                            table_md = self._table_to_markdown(table)
                            content_parts.append(table_md)
                            break
            
            content = '\n\n'.join(content_parts)
            
            metadata.update({
                "paragraph_count": len(doc.paragraphs),
                "table_count": table_count,
                "content_length": len(content),
            })
            
            content = self._sanitize_content(content)
            
            logger.info(
                f"Word 解析完成: {len(doc.paragraphs)} 段落, "
                f"{table_count} 个表格, "
                f"{len(content)} 字符"
            )
            
            return ParseResult(
                content=content,
                metadata=metadata,
                warnings=warnings
            )
            
        except ImportError:
            warnings.append("python-docx 未安装，Word 解析不可用")
            return ParseResult(
                content="",
                metadata=metadata,
                warnings=warnings
            )
        except Exception as e:
            warnings.append(f"Word 解析失败: {e}")
            return ParseResult(
                content="",
                metadata=metadata,
                warnings=warnings
            )
    
    def _table_to_markdown(self, table) -> str:
        """将 python-docx 表格转换为 Markdown
        
        Args:
            table: python-docx 表格对象
            
        Returns:
            str: Markdown 表格格式
        """
        rows_data = []
        
        for i, row in enumerate(table.rows):
            cells = [cell.text.strip() for cell in row.cells]
            rows_data.append(cells)
        
        return self._convert_table_to_markdown(rows_data)
    
    def _sanitize_content(self, content: str) -> str:
        """清理文档内容
        
        Args:
            content: 原始内容
            
        Returns:
            str: 清理后的内容
        """
        if not content:
            return ""
        
        # 移除多余的空白字符
        lines = []
        for line in content.split('\n'):
            lines.append(line.rstrip())
        
        content = '\n'.join(lines)
        
        # 移除连续的空行（最多保留两个）
        while '\n\n\n' in content:
            content = content.replace('\n\n\n', '\n\n')
        
        return content.strip()
    
    def is_available(self) -> bool:
        """检查解析器是否可用
        
        Returns:
            bool: 是否可用
        """
        return self._initialized and _magic_pdf_available
    
    def get_status(self) -> Dict[str, Any]:
        """获取解析器状态
        
        Returns:
            Dict: 状态信息
        """
        return {
            "available": self.is_available(),
            "gpu_available": self._has_gpu,
            "magic_pdf_available": _magic_pdf_available,
            "magic_pdf_error": _magic_pdf_import_error,
            "config": self.config.to_dict(),
        }
