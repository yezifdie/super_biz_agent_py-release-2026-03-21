"""向量索引服务模块

将本地文件索引到向量数据库的服务。
这是 RAG 系统的入库端：
1. 读取文件内容（支持 .txt、.md、.pdf、.docx 等）
2. 分割文档为小块（chunk）
3. 将 chunk 添加到 Milvus

索引流程：
目录 → 文件 → 智能解析（支持 MinerU）→ 分割文档 → 添加到向量存储

核心功能：
- index_directory(): 索引目录下所有支持的文件
- index_single_file(): 索引单个文件
- MinerU 智能解析：PDF、Word 文档的版面分析和表格识别
"""

from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Union

from loguru import logger

from app.services.document_splitter_service import document_splitter_service
from app.services.vector_store_manager import vector_store_manager

# 尝试导入文档解析器
try:
    from app.services.document_parser import create_parser
    _parser_factory_available = True
except ImportError:
    _parser_factory_available = False
    logger.warning("文档解析器模块不可用，仅支持 .txt 和 .md 文件")


class IndexingResult:
    """索引结果数据类

    封装索引操作的执行结果，包含：
    - 成功/失败状态
    - 文件统计（总数、成功数、失败数）
    - 耗时信息
    - 失败文件列表及错误原因
    """

    def __init__(self):
        """初始化索引结果（默认所有字段为空/零）"""
        self.success = False              # 是否全部成功
        self.directory_path = ""          # 被索引的目录路径
        self.total_files = 0              # 文件总数
        self.success_count = 0            # 成功数量
        self.fail_count = 0               # 失败数量
        self.start_time: Optional[datetime] = None   # 开始时间
        self.end_time: Optional[datetime] = None     # 结束时间
        self.error_message = ""           # 错误信息（如果有）
        self.failed_files: Dict[str, str] = {}  # 失败文件映射：路径 → 错误原因

    def increment_success_count(self):
        """增加成功计数"""
        self.success_count += 1

    def increment_fail_count(self):
        """增加失败计数"""
        self.fail_count += 1

    def add_failed_file(self, file_path: str, error: str):
        """添加失败文件记录

        Args:
            file_path: 失败的文件路径
            error: 错误原因描述
        """
        self.failed_files[file_path] = error

    def get_duration_ms(self) -> int:
        """获取索引耗时（毫秒）

        Returns:
            int: 耗时毫秒数，如果未完成则返回 0
        """
        if self.start_time and self.end_time:
            return int((self.end_time - self.start_time).total_seconds() * 1000)
        return 0

    def to_dict(self) -> Dict[str, Any]:
        """转换为字典格式

        Returns:
            可序列化为 JSON 的字典，包含所有索引统计信息
        """
        return {
            "success": self.success,
            "directory_path": self.directory_path,
            "total_files": self.total_files,
            "success_count": self.success_count,
            "fail_count": self.fail_count,
            "duration_ms": self.get_duration_ms(),
            "error_message": self.error_message,
            "failed_files": self.failed_files,
        }


class VectorIndexService:
    """向量索引服务

    负责将文件系统中的文档索引到 Milvus 向量数据库。

    支持的文件格式：
    - .txt: 纯文本文件
    - .md: Markdown 文件（会保留标题结构进行智能分割）
    - .pdf: PDF 文档（需要安装 magic-pdf 或使用降级解析）
    - .docx/.doc: Word 文档（需要安装 python-docx）

    索引策略：
    1. 使用智能解析器读取文件（MinerU > 降级解析器）
    2. 删除该文件的旧数据（如果存在）
    3. 分割文档为小块
    4. 添加到向量存储（自动生成向量）
    """
    
    # 支持的文件扩展名
    SUPPORTED_EXTENSIONS = ['.txt', '.md', '.markdown', '.pdf', '.docx', '.doc']

    def __init__(self):
        """初始化向量索引服务"""
        # 默认上传目录（可通过 index_directory() 参数覆盖）
        self.upload_path = "./uploads"
        # 文档解析器
        self._create_parser = create_parser if _parser_factory_available else None
        logger.info("向量索引服务初始化完成")

    def index_directory(self, directory_path: Optional[str] = None) -> IndexingResult:
        """索引指定目录下的所有支持的文件

        Args:
            directory_path: 要索引的目录路径
                           如果为 None，使用默认的 upload_path

        Returns:
            IndexingResult: 索引结果，包含成功/失败统计和详细信息

        支持的文件：
        - *.txt: 纯文本文件
        - *.md: Markdown 文件
        """
        result = IndexingResult()
        result.start_time = datetime.now()

        try:
            # 确定目标目录
            target_path = directory_path if directory_path else self.upload_path
            dir_path = Path(target_path).resolve()

            # 验证目录有效性
            if not dir_path.exists() or not dir_path.is_dir():
                raise ValueError(f"目录不存在或不是有效目录: {target_path}")

            result.directory_path = str(dir_path)

            # 查找支持的文件（包括 PDF 和 Word）
            files = []
            for ext in self.SUPPORTED_EXTENSIONS:
                files.extend(list(dir_path.glob(f"*{ext}")))

            # 无文件时的处理
            if not files:
                logger.warning(f"目录中没有找到支持的文件: {target_path}")
                result.total_files = 0
                result.success = True  # 没有文件不算失败
                result.end_time = datetime.now()
                return result

            result.total_files = len(files)
            logger.info(f"开始索引目录: {target_path}, 找到 {len(files)} 个文件")

            # 遍历索引每个文件
            for file_path in files:
                try:
                    self.index_single_file(str(file_path))
                    result.increment_success_count()
                    logger.info(f"✓ 文件索引成功: {file_path.name}")
                except Exception as e:
                    result.increment_fail_count()
                    result.add_failed_file(str(file_path), str(e))
                    logger.error(f"✗ 文件索引失败: {file_path.name}, 错误: {e}")

            # 判断是否全部成功
            result.success = result.fail_count == 0
            result.end_time = datetime.now()

            logger.info(
                f"目录索引完成: 总数={result.total_files}, "
                f"成功={result.success_count}, 失败={result.fail_count}"
            )

            return result

        except Exception as e:
            logger.error(f"索引目录失败: {e}")
            result.success = False
            result.error_message = str(e)
            result.end_time = datetime.now()
            return result

    def index_single_file(
        self, 
        file_path: str,
        parser_config: Optional[Dict[str, Any]] = None
    ):
        """索引单个文件

        索引流程：
        1. 使用智能解析器读取文件（MinerU > 降级解析器）
        2. 删除该文件的旧数据（幂等性：支持重复索引）
        3. 分割文档为小块
        4. 添加到向量存储

        Args:
            file_path: 文件的绝对路径或相对路径
            parser_config: 解析器配置（可选）
                - parse_mode: 解析模式 ("auto", "ocr", "txt")
                - ocr_enabled: 是否启用 OCR
                - table_mode: 表格解析模式 ("smart", "html", "markdown")

        Raises:
            ValueError: 文件不存在或不支持时抛出
            RuntimeError: 索引过程中发生错误时抛出
        """
        path = Path(file_path).resolve()

        if not path.exists() or not path.is_file():
            raise ValueError(f"文件不存在: {file_path}")

        ext = path.suffix.lower()
        if ext not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(f"不支持的文件类型: {ext}")

        logger.info(f"开始索引文件: {path}")

        try:
            # 步骤 1：使用智能解析器读取文件内容
            content = self._parse_file_content(path, parser_config)
            logger.info(f"文件解析完成: {path}, 内容长度: {len(content)} 字符")

            # 步骤 2：删除该文件的旧数据
            # 路径标准化为 POSIX 格式（使用正斜杠），确保一致性
            normalized_path = path.as_posix()
            vector_store_manager.delete_by_source(normalized_path)

            # 步骤 3：分割文档
            # document_splitter_service 会根据文件类型选择合适的分割策略
            documents = document_splitter_service.split_document(content, normalized_path)
            logger.info(f"文档分割完成: {file_path} -> {len(documents)} 个分片")

            # 步骤 4：添加到向量存储
            if documents:
                # vector_store_manager.add_documents 会自动：
                # 1. 为每个文档生成唯一 ID
                # 2. 调用 embedding_service 将文本转换为向量
                # 3. 批量插入 Milvus
                vector_store_manager.add_documents(documents)
                logger.info(f"文件索引完成: {file_path}, 共 {len(documents)} 个分片")
            else:
                logger.warning(f"文件内容为空或无法分割: {file_path}")

        except ValueError:
            raise
        except Exception as e:
            logger.error(f"索引文件失败: {file_path}, 错误: {e}")
            raise RuntimeError(f"索引文件失败: {e}") from e

    def _parse_file_content(
        self, 
        path: Path,
        parser_config: Optional[Dict[str, Any]] = None
    ) -> str:
        """使用智能解析器读取文件内容

        支持的解析策略：
        1. MinerU 解析器：PDF/Word 的智能解析（版面分析、表格识别）
        2. Markdown 解析器：保留标题结构
        3. 纯文本解析器：基础文本提取

        Args:
            path: 文件路径
            parser_config: 解析器配置

        Returns:
            str: 解析后的文本内容（Markdown 格式）
        """
        parser_config = parser_config or {}
        ext = path.suffix.lower()

        # 优先使用 ParserFactory 创建解析器
        if self._create_parser:
            try:
                parser = self._create_parser(str(path))
                result = parser.parse(str(path))

                # 记录解析警告
                if result.warnings:
                    for warning in result.warnings:
                        logger.warning(f"[{path.name}] {warning}")

                return result.content

            except Exception as e:
                logger.warning(f"ParserFactory 解析失败: {e}，尝试备用方法")
                # 继续使用备用方法

        # 备用方法：直接读取文本
        if ext in ['.txt', '.text']:
            return path.read_text(encoding="utf-8")
        elif ext in ['.md', '.markdown']:
            return path.read_text(encoding="utf-8")
        elif ext == '.pdf':
            return self._parse_pdf_fallback(path, parser_config)
        elif ext in ['.docx', '.doc']:
            return self._parse_docx_fallback(path)
        else:
            # 尝试 UTF-8 读取
            try:
                return path.read_text(encoding="utf-8")
            except UnicodeDecodeError:
                return path.read_text(encoding="gbk", errors="ignore")

    def _parse_pdf_fallback(
        self, 
        path: Path,
        parser_config: Optional[Dict[str, Any]] = None
    ) -> str:
        """PDF 降级解析（使用 PyMuPDF）

        Args:
            path: PDF 文件路径
            parser_config: 解析配置

        Returns:
            str: 提取的文本内容
        """
        try:
            import fitz  # PyMuPDF

            doc = fitz.open(str(path))
            content_parts = []

            for page_num in range(len(doc)):
                page = doc[page_num]
                text = page.get_text("text")

                if text.strip():
                    content_parts.append(f"## 第 {page_num + 1} 页\n\n{text.strip()}")

            doc.close()
            return '\n\n'.join(content_parts)

        except ImportError:
            logger.warning("PyMuPDF 未安装，PDF 解析不可用")
            return ""
        except Exception as e:
            logger.warning(f"PDF 解析失败: {e}")
            return ""

    def _parse_docx_fallback(self, path: Path) -> str:
        """Word 文档降级解析

        Args:
            path: Word 文件路径

        Returns:
            str: 提取的文本内容
        """
        try:
            from docx import Document

            doc = Document(str(path))
            content_parts = []

            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    # 判断是否是标题
                    if para.style.name.startswith('Heading'):
                        level = para.style.name.replace('Heading ', '')
                        try:
                            level = int(level)
                            content_parts.append(f"{'#' * level} {text}")
                        except ValueError:
                            content_parts.append(text)
                    else:
                        content_parts.append(text)

            # 处理表格
            for table in doc.tables:
                table_md = self._table_to_markdown(table)
                if table_md:
                    content_parts.append(table_md)

            return '\n\n'.join(content_parts)

        except ImportError:
            logger.warning("python-docx 未安装，Word 解析不可用")
            return ""
        except Exception as e:
            logger.warning(f"Word 解析失败: {e}")
            return ""

    def _table_to_markdown(self, table) -> str:
        """将 python-docx 表格转换为 Markdown

        Args:
            table: python-docx 表格对象

        Returns:
            str: Markdown 表格格式
        """
        lines = []

        for i, row in enumerate(table.rows):
            cells = [cell.text.strip().replace('|', '\\|') for cell in row.cells]
            lines.append('| ' + ' | '.join(cells) + ' |')

            if i == 0:
                separator = '| ' + ' | '.join(['---'] * len(cells)) + ' |'
                lines.append(separator)

        return '\n'.join(lines) if lines else ""

    def get_parser_status(self) -> Dict[str, Any]:
        """获取解析器状态

        Returns:
            Dict: 解析器状态信息
        """
        status = {
            "parser_factory_available": _parser_factory_available,
            "supported_extensions": self.SUPPORTED_EXTENSIONS,
        }

        if _parser_factory_available and self._parser_factory:
            status["supported_by_factory"] = self._parser_factory.get_supported_extensions()

            # MinerU 状态
            try:
                from app.services.document_parser import MinerUParser
                mineru = MinerUParser()
                status["mineru"] = mineru.get_status()
            except Exception as e:
                status["mineru"] = {"error": str(e)}

        return status


# ========== 全局向量索引服务单例 ==========
vector_index_service = VectorIndexService()
